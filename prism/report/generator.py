"""Report generator — produces Markdown, DOCX, and JSON outputs."""

from __future__ import annotations

import json
from datetime import datetime
from io import BytesIO
from typing import Any, Literal, Optional

from prism.models import AITrafficRecord, AIMetadata


# ════════════════════════════════════════════════════════════════════
# Markdown / text report
# ════════════════════════════════════════════════════════════════════

def _header_banner(task_id: str, total: int, generated_at: datetime) -> str:
    ts = generated_at.strftime("%Y-%m-%dT%H:%M:%SZ")
    return (
        "╔══════════════════════════════════════════════════════════════════╗\n"
        "║  PRISM AI Traffic Analysis Report                               ║\n"
        f"║  Task ID: {task_id:<54}║\n"
        f"║  Generated: {ts:<51}║\n"
        f"║  Total AI Requests Captured: {total:<35}║\n"
        "╚══════════════════════════════════════════════════════════════════╝"
    )


def _record_separator(seq: int, record: AITrafficRecord) -> str:
    label = f"Record #{seq} — {record.provider} {record.service_type.capitalize()}"
    model_str = f" | Model: {record.metadata.model_name}" if record.metadata.model_name else ""
    type_str = f"SSE Stream" if record.stream_type == "sse" else (
        "WebSocket" if record.stream_type == "websocket" else f"HTTP {record.method}"
    )
    return (
        "\n" + "═" * 66 + "\n"
        f" {label}\n"
        f" Provider: {record.provider}{model_str}\n"
        f" Type: {type_str} | Confidence: {record.confidence:.2f}\n"
        + "═" * 66
    )


def _metadata_block(meta: AIMetadata) -> str:
    lines = ["── Extracted Metadata " + "─" * 44]
    lines.append(f"  Provider     : {meta.provider}")
    if meta.model_name:
        lines.append(f"  Model        : {meta.model_name}")
    if meta.api_version:
        lines.append(f"  API Version  : {meta.api_version}")
    lines.append(f"  User-Agent   : {meta.user_agent or '(none)'}")
    if meta.client_version:
        lines.append(f"  Client Ver.  : {meta.client_version}")
    if meta.app_version:
        lines.append(f"  App Ver.     : {meta.app_version}")
    if meta.platform:
        lines.append(f"  Platform     : {meta.platform}")
    lines.append(f"  Auth Type    : {meta.auth_type}")
    lines.append(f"  Streaming    : {'Yes (' + (meta.is_streaming and 'SSE' or '') + ')' if meta.is_streaming else 'No'}")
    if meta.thinking_enabled is not None:
        lines.append(f"  Thinking     : {'Enabled' if meta.thinking_enabled else 'Disabled'}")
    if meta.file_upload:
        fu = meta.file_upload
        lines.append(f"  File Name    : {fu.filename}")
        lines.append(f"  File Size    : {fu.size:,} bytes")
    if meta.custom_headers:
        lines.append("  Custom Headers:")
        for k, v in meta.custom_headers.items():
            lines.append(f"    {k}: {v}")
    return "\n".join(lines)


def _sse_block(record: AITrafficRecord) -> str:
    if record.stream_type != "sse":
        return ""
    events = record.stream_events or []
    lines = [f"\n── SSE Response Stream ({len(events)} events) ──"]
    for ev in events:
        lines.append(f"data: {ev.get('data', '')}")
    lines.append("")
    lines.append("── Aggregated Response ──")
    lines.append(record.aggregated_response or "(empty)")
    return "\n".join(lines)


def _ws_block(record: AITrafficRecord) -> str:
    if record.stream_type != "websocket":
        return ""
    frames = record.stream_events or []
    lines = ["── WebSocket Frames ──"]
    for i, frame in enumerate(frames, 1):
        direction = "Client → Server" if frame.get("direction") == "client" else "Server → Client"
        ts = frame.get("timestamp_start", "")
        lines.append(f"--- Frame #{i} ({direction}) [{ts}] ---")
        lines.append(frame.get("payload", ""))
        lines.append("")
    return "\n".join(lines)


def generate_markdown(
    task_id: str,
    records: list[AITrafficRecord],
    generated_at: Optional[datetime] = None,
) -> str:
    if generated_at is None:
        generated_at = datetime.utcnow()

    parts = [_header_banner(task_id, len(records), generated_at)]

    for record in records:
        parts.append(_record_separator(record.sequence, record))
        parts.append("")
        parts.append(record.raw_request or "(no request data)")
        parts.append("")
        if record.stream_type == "sse":
            parts.append(_sse_block(record))
        elif record.stream_type == "websocket":
            parts.append(_ws_block(record))
        else:
            if record.raw_response:
                parts.append("── Response ──────────────────────────────────────────────────")
                parts.append(record.raw_response)
        parts.append("")
        parts.append(_metadata_block(record.metadata))
        parts.append("")

    return "\n".join(parts)


# ════════════════════════════════════════════════════════════════════
# JSON report
# ════════════════════════════════════════════════════════════════════

def generate_json(
    task_id: str,
    records: list[AITrafficRecord],
    generated_at: Optional[datetime] = None,
) -> str:
    if generated_at is None:
        generated_at = datetime.utcnow()

    payload = {
        "task_id": task_id,
        "generated_at": generated_at.isoformat(),
        "total_records": len(records),
        "records": [_record_to_dict(r) for r in records],
    }
    return json.dumps(payload, indent=2, default=str)


def _record_to_dict(r: AITrafficRecord) -> dict:
    return {
        "id": r.id,
        "sequence": r.sequence,
        "provider": r.provider,
        "service_type": r.service_type,
        "confidence": r.confidence,
        "method": r.method,
        "url": r.url,
        "response_status": r.response_status,
        "stream_type": r.stream_type,
        "aggregated_response": r.aggregated_response,
        "metadata": r.metadata.model_dump(),
        "raw_request": r.raw_request,
        "raw_response": r.raw_response,
        "stream_events": r.stream_events,
        "timestamp": r.timestamp.isoformat() if r.timestamp else None,
    }


# ════════════════════════════════════════════════════════════════════
# DOCX report
# ════════════════════════════════════════════════════════════════════

def generate_docx(
    task_id: str,
    records: list[AITrafficRecord],
    generated_at: Optional[datetime] = None,
) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is required: pip install python-docx")

    if generated_at is None:
        generated_at = datetime.utcnow()

    doc = Document()

    # Title
    title = doc.add_heading("PRISM AI Traffic Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Summary table
    doc.add_paragraph(f"Task ID: {task_id}")
    doc.add_paragraph(f"Generated: {generated_at.strftime('%Y-%m-%dT%H:%M:%SZ')}")
    doc.add_paragraph(f"Total AI Requests Captured: {len(records)}")
    doc.add_paragraph("")

    for record in records:
        # Record heading
        heading_text = (
            f"Record #{record.sequence} — {record.provider} "
            f"{record.service_type.capitalize()}"
        )
        doc.add_heading(heading_text, level=2)

        # Summary
        p = doc.add_paragraph()
        p.add_run("Provider: ").bold = True
        p.add_run(record.provider)
        if record.metadata.model_name:
            p.add_run(" | Model: ").bold = True
            p.add_run(record.metadata.model_name)
        p.add_run(" | Confidence: ").bold = True
        p.add_run(f"{record.confidence:.2f}")

        # Raw request
        doc.add_heading("Raw Request", level=3)
        req_para = doc.add_paragraph(record.raw_request or "")
        req_para.style = doc.styles["No Spacing"]
        for run in req_para.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

        # Response / SSE
        if record.stream_type == "sse":
            doc.add_heading("SSE Stream", level=3)
            events = record.stream_events or []
            doc.add_paragraph(f"Events: {len(events)}")
            doc.add_heading("Aggregated Response", level=4)
            doc.add_paragraph(record.aggregated_response or "(empty)")
        elif record.raw_response:
            doc.add_heading("Raw Response", level=3)
            resp_para = doc.add_paragraph(record.raw_response)
            resp_para.style = doc.styles["No Spacing"]
            for run in resp_para.runs:
                run.font.name = "Courier New"
                run.font.size = Pt(8)

        # Metadata table
        doc.add_heading("Metadata", level=3)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"

        meta = record.metadata
        rows_data = [
            ("Provider", meta.provider),
            ("Model", meta.model_name or ""),
            ("API Version", meta.api_version or ""),
            ("User-Agent", meta.user_agent),
            ("Client Version", meta.client_version or ""),
            ("Platform", meta.platform or ""),
            ("Auth Type", meta.auth_type),
            ("Streaming", "Yes" if meta.is_streaming else "No"),
        ]
        if meta.file_upload:
            rows_data.append(("File Name", meta.file_upload.filename))
            rows_data.append(("File Size", f"{meta.file_upload.size:,} bytes"))

        for field, value in rows_data:
            row = table.add_row().cells
            row[0].text = field
            row[1].text = value

        doc.add_paragraph("")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════════════
# Dispatcher
# ════════════════════════════════════════════════════════════════════

def generate_report(
    task_id: str,
    records: list[AITrafficRecord],
    fmt: Literal["markdown", "docx", "json"] = "markdown",
    generated_at: Optional[datetime] = None,
) -> tuple[bytes, str]:
    """
    Returns (content_bytes, content_type).
    """
    if fmt == "json":
        text = generate_json(task_id, records, generated_at)
        return text.encode("utf-8"), "application/json"
    if fmt == "docx":
        data = generate_docx(task_id, records, generated_at)
        return data, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    text = generate_markdown(task_id, records, generated_at)
    return text.encode("utf-8"), "text/markdown; charset=utf-8"
